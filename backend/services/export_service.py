from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import pandas as pd
from sqlalchemy.orm import Session, joinedload
from models.models import Enseignant, Examen, Affectation
from datetime import datetime, date
from typing import List, Dict
import os
import sys
from config import EXPORT_DIR
import logging

logger = logging.getLogger(__name__)


def get_resource_path(relative_path):
    """
    Obtient le chemin absolu d'une ressource, que ce soit en mode développement ou PyInstaller.
    
    Args:
        relative_path: Chemin relatif de la ressource
        
    Returns:
        Chemin absolu de la ressource
    """
    try:
        # PyInstaller crée un dossier temporaire et stocke le chemin dans _MEIPASS
        base_path = sys._MEIPASS
    except AttributeError:
        # En mode développement, utilise le répertoire courant
        base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    return os.path.join(base_path, relative_path)


class ExportService:
    """Service pour l'export de documents (PDF, Word, Excel)"""
    
    def __init__(self, db: Session):
        self.db = db
        self.export_dir = EXPORT_DIR
    
    @staticmethod
    def _convertir_session(code_session: str) -> str:
        """Convertit le code de session en texte complet
        
        Args:
            code_session: Code de la session (P, R, etc.)
            
        Returns:
            Texte complet de la session (Principale, Rattrapage, etc.)
        """
        conversions = {
            'P': 'Principale',
            'R': 'Rattrapage',
            'C': 'Contrôle'  # Au cas où
        }
        return conversions.get(code_session, code_session)
    
    @staticmethod
    def _determiner_numero_seance(h_debut) -> str:
        """Détermine le numéro de séance en fonction de l'heure de début
        
        Args:
            h_debut: Heure de début de la séance (time object ou string "HH:MM")
            
        Returns:
            Numéro de séance (S1, S2, S3, S4)
        """
        from datetime import time
        
        # Si c'est un string, le convertir en time object
        if isinstance(h_debut, str):
            heures, minutes = map(int, h_debut.split(':'))
            h_debut = time(heures, minutes)
        
        # Convertir en minutes depuis minuit pour faciliter la comparaison
        heure_minutes = h_debut.hour * 60 + h_debut.minute
        
        # Définir les plages horaires
        # S1: 8:30-10:00 (510-600 minutes)
        # S2: 10:30-12:00 (630-720 minutes)
        # S3: 12:30-14:00 (750-840 minutes)  
        # S4: 14:30-16:00 (870-960 minutes)
        
        if 510 <= heure_minutes < 630:  # 8:30 - 10:29
            return "S1"
        elif 630 <= heure_minutes < 750:  # 10:30 - 12:29
            return "S2"
        elif 750 <= heure_minutes < 870:  # 12:30 - 14:29
            return "S3"
        elif 870 <= heure_minutes < 1020:  # 14:30 - 16:59
            return "S4"
        else:
            # Par défaut, retourner basé sur l'heure
            if heure_minutes < 630:
                return "S1"
            elif heure_minutes < 750:
                return "S2"
            elif heure_minutes < 870:
                return "S3"
            else:
                return "S4"
    
    @staticmethod
    def _obtenir_horaires_seance(seance: str) -> tuple:
        """Retourne les horaires (h_debut, h_fin) pour un numéro de séance
        
        Args:
            seance: Numéro de séance (S1, S2, S3, S4)
            
        Returns:
            Tuple (h_debut, h_fin) au format string "HH:MM"
        """
        horaires = {
            'S1': ('08:30', '10:00'),
            'S2': ('10:30', '12:00'),
            'S3': ('12:30', '14:00'),
            'S4': ('14:30', '16:00')
        }
        
        seance_upper = seance.upper()
        if seance_upper not in horaires:
            raise ValueError(f"Séance invalide '{seance}'. Doit être S1, S2, S3 ou S4")
        
        return horaires[seance_upper]
    
    def generer_planning_global_pdf(self, date_debut: date = None, date_fin: date = None) -> str:
        """Génère le planning global en PDF"""
        filename = f"planning_global_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(self.export_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        elements = []
        styles = getSampleStyleSheet()
        
        # Titre
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#1a237e'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        elements.append(Paragraph("PLANNING GLOBAL DES SURVEILLANCES", title_style))
        elements.append(Spacer(1, 0.5*cm))
        
        # Requête des examens avec filtres de date
        query = self.db.query(Examen)
        if date_debut:
            query = query.filter(Examen.dateExam >= date_debut)
        if date_fin:
            query = query.filter(Examen.dateExam <= date_fin)
        
        examens = query.order_by(Examen.dateExam, Examen.h_debut).all()
        
        # Grouper par date
        examens_par_date = {}
        for examen in examens:
            date_key = examen.dateExam
            if date_key not in examens_par_date:
                examens_par_date[date_key] = []
            examens_par_date[date_key].append(examen)
        
        # Générer le contenu pour chaque date
        for date_exam, liste_examens in sorted(examens_par_date.items()):
            # Titre de la date
            date_style = ParagraphStyle(
                'DateStyle',
                parent=styles['Heading2'],
                fontSize=12,
                textColor=colors.HexColor('#0d47a1'),
                spaceAfter=10
            )
            elements.append(Paragraph(f"Date: {date_exam.strftime('%d/%m/%Y')}", date_style))
            
            # Tableau des examens
            data = [['Horaire', 'Salle', 'Session', 'Semestre', 'Surveillants']]
            
            for examen in liste_examens:
                affectations = self.db.query(Affectation).options(
                    joinedload(Affectation.enseignant)
                ).filter(
                    Affectation.examen_id == examen.id
                ).all()
                
                surveillants = []
                for aff in affectations:
                    ens = aff.enseignant
                    nom_complet = f"{ens.nom} {ens.prenom}"
                    if aff.est_responsable:
                        nom_complet += " (R)"
                    surveillants.append(nom_complet)
                
                # Enlever les doublons (car plusieurs enseignants peuvent être dans la même salle)
                surveillants = list(dict.fromkeys(surveillants))
                
                horaire = f"{examen.h_debut.strftime('%H:%M')} - {examen.h_fin.strftime('%H:%M')}"
                salle = examen.cod_salle
                # Convertir le code de session en texte complet
                session = self._convertir_session(examen.session)
                semestre = examen.semestre
                surveillants_str = "\n".join(surveillants[:3]) if surveillants else "Non affecté"
                if len(surveillants) > 3:
                    surveillants_str += f"\n+ {len(surveillants) - 3} autres"
                
                data.append([horaire, salle, session, semestre, surveillants_str])
            
            table = Table(data, colWidths=[3*cm, 2.5*cm, 2*cm, 3*cm, 7*cm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            
            elements.append(table)
            elements.append(Spacer(1, 0.5*cm))
        
        # Générer le PDF
        doc.build(elements)
        logger.info(f"✅ Planning global PDF généré: {filepath}")
        return filepath
    
    def generer_convocations_individuelles(self) -> List[str]:
        """Génère les convocations individuelles pour chaque enseignant (Word + PDF)"""
        filepaths = []
        
        enseignants = self.db.query(Enseignant).all()
        
        for enseignant in enseignants:
            # Récupérer les affectations
            affectations = self.db.query(Affectation).options(
                joinedload(Affectation.examen)
            ).filter(
                Affectation.enseignant_id == enseignant.id
            ).join(Examen).order_by(Examen.dateExam, Examen.h_debut).all()
            
            if not affectations:
                continue
            
            # Générer Word
            filename_word = f"convocation_{enseignant.nom}_{enseignant.prenom}_{datetime.now().strftime('%Y%m%d')}.docx"
            filepath_word = os.path.join(self.export_dir, filename_word)
            self._generer_convocation_word(enseignant, affectations, filepath_word)
            filepaths.append(filepath_word)
        
        logger.info(f"✅ {len(filepaths)} convocations générées")
        return filepaths
    
    def _generer_convocation_word(self, enseignant: Enseignant, affectations: List[Affectation], filepath: str):
        """Génère une convocation Word pour un enseignant"""
        doc = Document()
        
        # Ajouter l'en-tête ISI
        self._ajouter_entete(doc)
        
        # Ajouter le pied de page ISI
        self._ajouter_pied_de_page(doc)
        
        # Ajouter un retour à la ligne après l'entête
        doc.add_paragraph()
        
        # En-tête "Note à" - centré et en bleu
        note_para = doc.add_paragraph()
        note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_note = note_para.add_run('Notes à\n')
        run_note.bold = True
        run_note.font.size = Pt(14)
        run_note.font.color.rgb = RGBColor(0, 51, 153)
        
        run_nom = note_para.add_run(f'M./Mme {enseignant.nom} {enseignant.prenom}')
        run_nom.font.size = Pt(12)
        run_nom.font.color.rgb = RGBColor(0, 51, 153)
        
        doc.add_paragraph()
        
        # Message d'accueil
        doc.add_paragraph('Cher(e) Collègue,')
        
        doc.add_paragraph()
        
        # Message principal
        doc.add_paragraph(
            "Vous êtes prié(e) d'assurer la surveillance et (ou) la responsabilité des examens "
            "selon le calendrier ci-joint."
        )
        
        doc.add_paragraph()
        
        # Regrouper les affectations par séance (date + horaires) pour éviter les doublons
        seances = {}
        for aff in affectations:
            examen = aff.examen
            key = (examen.dateExam, examen.h_debut, examen.h_fin)
            if key not in seances:
                seances[key] = {
                    'date': examen.dateExam,
                    'h_debut': examen.h_debut,
                    'h_fin': examen.h_fin
                }
        
        # Trier par date puis heure de début
        seances_list = sorted(seances.values(), key=lambda x: (x['date'], x['h_debut']))
        
        # Tableau des surveillances simplifié : Date | Heure | Durée
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Définir les largeurs de colonnes
        table.autofit = False
        table.allow_autofit = False
        for row in table.rows:
            row.cells[0].width = Inches(2.5)   # Date
            row.cells[1].width = Inches(2.0)   # Heure
            row.cells[2].width = Inches(2.0)   # Durée
        
        hdr_cells = table.rows[0].cells
        
        # Styliser l'en-tête du tableau
        for i, header_text in enumerate(['Date', 'Heure', 'Durée']):
            cell = hdr_cells[i]
            cell.paragraphs[0].clear()
            self._set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(255, 255, 255)  # Texte blanc
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Fond bleu vif pour l'en-tête
            self._set_cell_background(cell, "003399")
        
        for seance in seances_list:
            row_cells = table.add_row().cells
            
            # Date
            self._set_cell_vertical_alignment(row_cells[0], 'center')
            row_cells[0].text = seance['date'].strftime('%d/%m/%Y')
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Heure
            self._set_cell_vertical_alignment(row_cells[1], 'center')
            row_cells[1].text = seance['h_debut'].strftime('%H:%M')
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Calculer la durée
            h_debut = seance['h_debut']
            h_fin = seance['h_fin']
            # Convertir en minutes pour calculer la durée
            debut_minutes = h_debut.hour * 60 + h_debut.minute
            fin_minutes = h_fin.hour * 60 + h_fin.minute
            duree_minutes = fin_minutes - debut_minutes
            
            # Formater la durée (ex: "2h00" ou "1h30")
            heures = duree_minutes // 60
            minutes = duree_minutes % 60
            if minutes > 0:
                duree_str = f"{heures}h{minutes:02d}"
            else:
                duree_str = f"{heures}h00"
            
            # Durée
            self._set_cell_vertical_alignment(row_cells[2], 'center')
            row_cells[2].text = duree_str
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Pied de page

        
        # Sauvegarder
        doc.save(filepath)
    
    def generer_listes_par_creneau(self) -> List[str]:
        """Génère les listes de surveillants par créneau (Word) - Un fichier par jour avec toutes les séances"""
        filepaths = []
        
        # Récupérer tous les examens
        examens = self.db.query(Examen).order_by(
            Examen.dateExam, Examen.h_debut
        ).all()
        
        # Grouper d'abord par date, puis par créneau
        jours = {}
        for examen in examens:
            date_key = examen.dateExam
            if date_key not in jours:
                jours[date_key] = {}
            
            creneau_key = (examen.h_debut, examen.h_fin)
            if creneau_key not in jours[date_key]:
                jours[date_key][creneau_key] = []
            jours[date_key][creneau_key].append(examen)
        
        # Générer un document par jour avec toutes ses séances
        for date_exam, creneaux in sorted(jours.items()):
            # Utiliser le même format que generer_liste_creneau_specifique: YYYY-MM-DD
            filename = f"liste_seance_{date_exam.strftime('%Y-%m-%d')}.docx"
            filepath = os.path.join(self.export_dir, filename)
            
            # Créer un document avec toutes les séances du jour
            doc = Document()
            
            # Ajouter l'en-tête au document
            self._ajouter_entete(doc)
            
            # Ajouter le pied de page au document
            self._ajouter_pied_de_page(doc)
            
            # Trier les créneaux par heure de début
            creneaux_tries = sorted(creneaux.items(), key=lambda x: x[0][0])
            
            for idx, ((h_debut, h_fin), liste_examens) in enumerate(creneaux_tries):
                # Ajouter le contenu de la séance
                self._ajouter_seance_au_document(doc, date_exam, h_debut, h_fin, liste_examens)
                
                # Ajouter un saut de page entre les séances (chaque séance sur sa propre page)
                if idx < len(creneaux_tries) - 1:
                    doc.add_page_break()
            
            # Sauvegarder le document
            doc.save(filepath)
            filepaths.append(filepath)
        
        logger.info(f"✅ {len(filepaths)} fichiers générés (un par jour)")
        return filepaths
    
    def _ajouter_seance_au_document(
        self, 
        doc: Document,
        date_exam: date, 
        h_debut, 
        h_fin, 
        examens: List[Examen]
    ):
        """Ajoute une séance à un document Word existant avec mise en forme professionnelle"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Pour chaque groupe (même session/semestre) dans ce créneau
        # Regrouper les examens par session et semestre
        groupes = {}
        for examen in examens:
            cle = (examen.session, examen.semestre)
            if cle not in groupes:
                groupes[cle] = []
            groupes[cle].append(examen)
        
        # Générer un tableau pour chaque groupe session/semestre
        for (session, semestre), examens_groupe in groupes.items():
            # Convertir le code de session en texte complet
            session_text = self._convertir_session(session)
            
            # Déterminer le numéro de séance en fonction de l'heure
            numero_seance = self._determiner_numero_seance(h_debut)
            # Ajout d'un retour à la ligne après l'entête
            doc.add_paragraph()
            # === INFORMATIONS DE LA SÉANCE (TITRE CENTRÉ ET EN GRAS) ===
            # Ligne 1 : AU, Semestre, Session - centré, gras, plus grand
            p_info = doc.add_paragraph()
            p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_info = p_info.add_run(
                f"AU : 2024-2025 - Semestre : {semestre.split()[-1]} - Session : {session_text}"
            )
            run_info.font.bold = True
            run_info.font.size = Pt(14)
            run_info.font.color.rgb = RGBColor(0, 51, 153)  # Bleu plus vif
            p_info.space_after = Pt(6)
            
            # Ligne 2 : Date et horaire - centré, gras, plus grand
            p_date = doc.add_paragraph()
            p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_date = p_date.add_run(
                f"Date : {date_exam.strftime('%d/%m/%Y')} - Séance : {numero_seance}"
            )
            run_date.font.bold = True
            run_date.font.size = Pt(12)
            run_date.font.color.rgb = RGBColor(0, 51, 153)  # Bleu plus vif
            p_date.space_after = Pt(12)
            
            # === TABLEAU DES SURVEILLANTS ===
            # Récupérer tous les enseignants affectés à ce créneau
            enseignants_affectes = []
            for examen in examens_groupe:
                affectations = self.db.query(Affectation).options(
                    joinedload(Affectation.enseignant)
                ).filter(
                    Affectation.examen_id == examen.id
                ).all()
                
                for aff in affectations:
                    # Éviter les doublons
                    if not any(e['id'] == aff.enseignant.id for e in enseignants_affectes):
                        enseignants_affectes.append({
                            'id': aff.enseignant.id,
                            'nom': aff.enseignant.nom,
                            'prenom': aff.enseignant.prenom
                        })
            
            # Créer le tableau avec 3 colonnes : Enseignant | Salle | Signature
            if enseignants_affectes:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Définir les largeurs de colonnes (ajustées pour tenir dans la page)
                table.autofit = False
                table.allow_autofit = False
                
                # Largeur totale disponible : environ 7 pouces (marges de 0.5" de chaque côté)
                for row in table.rows:
                    row.cells[0].width = Inches(3.5)   # Enseignant (50%)
                    row.cells[1].width = Inches(1.75)  # Salle (25%)
                    row.cells[2].width = Inches(1.75)  # Signature (25%)
                
                # Définir la largeur totale du tableau
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tbl = table._element
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                
                # Définir la largeur du tableau en pourcentage
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), '5000')  # 100% de la largeur disponible
                tblW.set(qn('w:type'), 'pct')
                tblPr.append(tblW)
                
                # En-têtes avec style
                hdr_cells = table.rows[0].cells
                for i, header_text in enumerate(['Enseignant', 'Salle', 'Signature']):
                    cell = hdr_cells[i]
                    cell.paragraphs[0].clear()
                    self._set_cell_vertical_alignment(cell, 'center')
                    p = cell.paragraphs[0]
                    run = p.add_run(header_text)
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Texte blanc
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Fond bleu vif pour l'en-tête
                    self._set_cell_background(cell, "003399")  # Bleu plus vif et saturé
                
                # Ajouter chaque enseignant
                for ens in sorted(enseignants_affectes, key=lambda x: (x['nom'], x['prenom'])):
                    row_cells = table.add_row().cells
                    
                    # Définir la hauteur de ligne compacte
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    
                    tr = row_cells[0]._element.getparent()
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    # Hauteur fixe de 300 twips (environ 0.53 cm) - compact mais lisible
                    trHeight.set(qn('w:val'), '300')
                    trHeight.set(qn('w:hRule'), 'exact')  # Hauteur exacte
                    trPr.append(trHeight)
                    
                    # Nom de l'enseignant
                    self._set_cell_vertical_alignment(row_cells[0], 'center')
                    row_cells[0].text = f"{ens['nom']} {ens['prenom']}"
                    row_cells[0].paragraphs[0].runs[0].font.size = Pt(8)  # Police plus petite
                    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Salle (vide)
                    self._set_cell_vertical_alignment(row_cells[1], 'center')
                    row_cells[1].text = ''
                    
                    # Signature (vide)
                    self._set_cell_vertical_alignment(row_cells[2], 'center')
                    row_cells[2].text = ''
                    
                    # Aucun espacement pour maximiser l'espace
                    for cell in row_cells:
                        cell.paragraphs[0].space_after = Pt(0)
                        cell.paragraphs[0].space_before = Pt(0)
            else:
                doc.add_paragraph("⚠️ Aucun surveillant affecté", style='Intense Quote')
            
            # Ne rien ajouter ici - le pied de page sera ajouté via la section footer du document
    
    def _set_table_borders(self, table):
        """Définit les bordures d'un tableau (style simple)"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # Taille de bordure
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Noir
            tblBorders.append(border)
        tblPr.append(tblBorders)
    
    def _set_cell_background(self, cell, color_hex):
        """Définit la couleur de fond d'une cellule"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        cell_properties = cell._element.get_or_add_tcPr()
        cell_shading = OxmlElement('w:shd')
        cell_shading.set(qn('w:fill'), color_hex)
        cell_properties.append(cell_shading)
    
    def _set_cell_vertical_alignment(self, cell, alignment='center'):
        """Définit l'alignement vertical d'une cellule
        
        Args:
            cell: Cellule du tableau
            alignment: 'top', 'center', 'bottom' (défaut: 'center')
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), alignment)
        tcPr.append(vAlign)
    
    def _add_horizontal_line(self, doc, color_hex="000000"):
        """Ajoute une ligne horizontale en utilisant une bordure de paragraphe
        
        Args:
            doc: Document Word
            color_hex: Couleur de la ligne en hexadécimal (sans #)
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        
        # Créer l'élément de bordure
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # Épaisseur
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        # Réduire l'espacement
        p.space_before = Pt(3)
        p.space_after = Pt(3)
    
    def _add_horizontal_line_to_section(self, section, color_hex="000000"):
        """Ajoute une ligne horizontale dans une section (header ou footer)
        
        Args:
            section: Section du document (header ou footer)
            color_hex: Couleur de la ligne en hexadécimal (sans #)
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        p = section.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        
        # Créer l'élément de bordure
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    
    def _ajouter_entete(self, doc):
        """Ajoute l'en-tête au document Word
        
        Args:
            doc: Document Word
        """
        # Accéder à la section du document
        section = doc.sections[0]
        header = section.header
        
        # Définir les marges de la section pour supprimer les espaces à gauche
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # Créer un tableau pour l'en-tête avec 3 lignes et 3 colonnes
        # La largeur totale doit être spécifiée pour les tableaux dans header/footer
        header_table = header.add_table(rows=3, cols=3, width=Inches(7.5))
        header_table.autofit = False
        header_table.allow_autofit = False
        
        # Définir les largeurs de colonnes
        for row in header_table.rows:
            row.cells[0].width = Inches(1.5)  # Logo (réduit)
            row.cells[1].width = Inches(4.5)  # Titres (augmenté)
            row.cells[2].width = Inches(1.5)  # Référence
        
        # === LIGNE 1 : Titre principal et référence ===
        # Logo ISI (ligne 1, colonne 1) - sera fusionné avec les lignes 2 et 3
        cell_logo_row1 = header_table.rows[0].cells[0]
        cell_logo_row1.text = ""
        
        # Titre principal (ligne 1, colonne 2)
        cell_titre1 = header_table.rows[0].cells[1]
        cell_titre1.text = ""
        self._set_cell_vertical_alignment(cell_titre1, 'center')
        p_titre1 = cell_titre1.paragraphs[0]
        run_titre1 = p_titre1.add_run("GESTION DES EXAMENS ET DÉLIBÉRATIONS")
        run_titre1.font.size = Pt(13)
        run_titre1.font.bold = True
        run_titre1.font.color.rgb = RGBColor(0, 51, 153)
        p_titre1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Référence document (ligne 1, colonne 3)
        cell_ref1 = header_table.rows[0].cells[2]
        cell_ref1.text = ""
        self._set_cell_vertical_alignment(cell_ref1, 'center')
        p_ref1 = cell_ref1.paragraphs[0]
        run_ref1 = p_ref1.add_run("EXD-FR-08-01")
        run_ref1.font.size = Pt(10)
        run_ref1.font.bold = True
        run_ref1.font.color.rgb = RGBColor(0, 51, 153)
        p_ref1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # === LIGNE 2 : Sous-titre et date d'approbation ===
        # Logo ISI (fusionner les 3 lignes pour la colonne 1)
        cell_logo = header_table.rows[0].cells[0]
        cell_logo.merge(header_table.rows[1].cells[0])
        cell_logo.merge(header_table.rows[2].cells[0])
        cell_logo.text = ""
        self._set_cell_vertical_alignment(cell_logo, 'center')
        
        # Ajouter l'image du logo ISI
        logo_path = get_resource_path(os.path.join('logo', 'logoISI.png'))
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.2))
        
        # Sous-titre (ligne 2, colonne 2)
        cell_titre2 = header_table.rows[1].cells[1]
        cell_titre2.text = ""
        self._set_cell_vertical_alignment(cell_titre2, 'center')
        p_titre2 = cell_titre2.paragraphs[0]
        run_titre2 = p_titre2.add_run("Procédure d'exécution des épreuves")
        run_titre2.font.size = Pt(13)
        run_titre2.font.color.rgb = RGBColor(0, 51, 153)
        run_titre2.font.bold = True
        p_titre2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date d'approbation (ligne 2, colonne 3)
        cell_ref2 = header_table.rows[1].cells[2]
        cell_ref2.text = ""
        self._set_cell_vertical_alignment(cell_ref2, 'center')
        p_ref2 = cell_ref2.paragraphs[0]
        run_ref2_label = p_ref2.add_run("Date d'approbation ")
        run_ref2_label.font.size = Pt(10)
        run_ref2_label.font.color.rgb = RGBColor(0, 51, 153)
        run_ref2_label.font.bold = True
        p_ref2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ref2_date = cell_ref2.add_paragraph()
        run_ref2_date = p_ref2_date.add_run("0504-24")
        run_ref2_date.font.size = Pt(10)
        run_ref2_date.font.bold = True
        run_ref2_date.font.color.rgb = RGBColor(0, 51, 153)
        p_ref2_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # === LIGNE 3 : Liste d'affectation et page ===
        # Liste d'affectation (ligne 3, colonne 2)
        cell_titre3 = header_table.rows[2].cells[1]
        cell_titre3.text = ""
        self._set_cell_vertical_alignment(cell_titre3, 'center')
        p_titre3 = cell_titre3.paragraphs[0]
        run_titre3 = p_titre3.add_run("Liste d'affectation des surveillants")
        run_titre3.font.size = Pt(12)
        run_titre3.font.bold = True
        run_titre3.font.color.rgb = RGBColor(0, 51, 153)
        p_titre3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page (ligne 3, colonne 3)
        cell_ref3 = header_table.rows[2].cells[2]
        cell_ref3.text = ""
        self._set_cell_vertical_alignment(cell_ref3, 'center')
        p_ref3 = cell_ref3.paragraphs[0]
        run_ref3 = p_ref3.add_run("Page 1/1")
        run_ref3.font.size = Pt(10)
        run_ref3.font.bold = True
        run_ref3.font.color.rgb = RGBColor(0, 51, 153)
        p_ref3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter les bordures au tableau d'en-tête
        self._set_table_borders(header_table)
        
    
    def _ajouter_pied_de_page(self, doc):
        """Ajoute le pied de page au document Word
        
        Args:
            doc: Document Word
        """
        # Accéder à la section du document
        section = doc.sections[0]
        footer = section.footer
        
        # Ligne de séparation bleue
        p_line = footer.paragraphs[0]
        pPr = p_line._element.get_or_add_pPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), '003399')  # Bleu plus vif
        pBdr.append(top)
        pPr.append(pBdr)
        
        # Adresse en arabe
        p_arabic = footer.add_paragraph()
        run_arabic = p_arabic.add_run(
" نهج 02 أبو الريحان البيروني 2080 أريانة الهاتف: 71706164 الفاكس: 71706698 البريد الإلكتروني"
        )
        run_arabic.font.size = Pt(8)
        
        # Email en bleu vif
        run_email1 = p_arabic.add_run(" ISI@isi.rnu.tn")
        run_email1.font.size = Pt(8)
        run_email1.font.color.rgb = RGBColor(0, 51, 204)  # Bleu plus vif
        run_email1.font.underline = True
        
        p_arabic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_arabic.paragraph_format.space_after = Pt(2)
        
        # Adresse en français
        p_french = footer.add_paragraph()
        run_french = p_french.add_run(
            "02 Rue Abou Raihane Bayrouni 2080 Ariana   Tél :71706164   Email : "
        )
        run_french.font.size = Pt(8)
        
        # Email en bleu vif
        run_email2 = p_french.add_run("ISI@isi.rnu.tn")
        run_email2.font.size = Pt(8)
        run_email2.font.color.rgb = RGBColor(0, 51, 204)  # Bleu plus vif
        run_email2.font.underline = True
        
        p_french.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_french.paragraph_format.space_after = Pt(2)
    
    def generer_excel_global(self, date_debut: date = None, date_fin: date = None) -> str:
        """Génère un fichier Excel avec toutes les affectations"""
        filename = f"planning_excel_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        filepath = os.path.join(self.export_dir, filename)
        
        # Récupérer les données
        query = self.db.query(Affectation).options(
            joinedload(Affectation.examen),
            joinedload(Affectation.enseignant)
        ).join(Examen)
        if date_debut:
            query = query.filter(Examen.dateExam >= date_debut)
        if date_fin:
            query = query.filter(Examen.dateExam <= date_fin)
        
        affectations = query.all()
        
        # Préparer les données
        data = []
        for aff in affectations:
            examen = aff.examen
            enseignant = aff.enseignant
            
            # Convertir le code de session en texte complet
            session_text = self._convertir_session(examen.session)
            
            data.append({
                'Date': examen.dateExam.strftime('%d/%m/%Y'),
                'Horaire Début': examen.h_debut.strftime('%H:%M'),
                'Horaire Fin': examen.h_fin.strftime('%H:%M'),
                'Salle': aff.cod_salle,
                'Session': session_text,
                'Type': examen.type_ex,
                'Semestre': examen.semestre,
                'Enseignant': f"{enseignant.nom} {enseignant.prenom}",
                'Grade': enseignant.grade_code,
                'Email': enseignant.email,
                'Rôle': "Responsable" if aff.est_responsable else "Surveillant"
            })
        
        # Créer le DataFrame
        df = pd.DataFrame(data)
        
        # Sauvegarder
        df.to_excel(filepath, index=False, sheet_name='Planning')
        
        logger.info(f"✅ Planning Excel généré: {filepath}")
        return filepath
    
    def generer_convocation_enseignant(self, enseignant_id: int) -> str:
        """Génère la convocation pour un enseignant spécifique
        
        Args:
            enseignant_id: ID de l'enseignant
            
        Returns:
            Chemin du fichier Word généré
        """
        # Récupérer l'enseignant
        enseignant = self.db.query(Enseignant).filter(Enseignant.id == enseignant_id).first()
        if not enseignant:
            raise ValueError(f"Enseignant avec l'ID {enseignant_id} introuvable")
        
        # Récupérer les affectations
        affectations = self.db.query(Affectation).options(
            joinedload(Affectation.examen)
        ).filter(
            Affectation.enseignant_id == enseignant.id
        ).join(Examen).order_by(Examen.dateExam, Examen.h_debut).all()
        
        if not affectations:
            raise ValueError(f"Aucune affectation trouvée pour l'enseignant {enseignant.nom} {enseignant.prenom}")
        
        # Générer Word
        filename_word = f"convocation_{enseignant.nom}_{enseignant.prenom}.docx"
        filepath_word = os.path.join(self.export_dir, filename_word)
        self._generer_convocation_word(enseignant, affectations, filepath_word)
        
        logger.info(f"✅ Convocation générée pour {enseignant.nom} {enseignant.prenom}: {filepath_word}")
        return filepath_word
    
    def generer_liste_creneau_specifique(self, date_exam: date, seance: str) -> str:
        """Génère la liste des surveillants pour un créneau spécifique
        
        Args:
            date_exam: Date de l'examen
            seance: Numéro de séance (S1, S2, S3, S4)
            
        Returns:
            Chemin du fichier Word généré
        """
        from datetime import time as dt_time
        
        # Obtenir les horaires de la séance (format string "HH:MM")
        h_debut_str, h_fin_str = self._obtenir_horaires_seance(seance)
        
        # Convertir les strings en objets time pour la requête SQL
        h_debut_time = dt_time(*map(int, h_debut_str.split(':')))
        h_fin_time = dt_time(*map(int, h_fin_str.split(':')))
        
        # Récupérer les examens pour ce créneau
        examens = self.db.query(Examen).filter(
            Examen.dateExam == date_exam,
            Examen.h_debut == h_debut_time,
            Examen.h_fin == h_fin_time
        ).all()
        
        if not examens:
            raise ValueError(f"Aucun examen trouvé pour la séance {seance.upper()} du {date_exam}")
        
        # Générer le nom du fichier
        filename = f"liste_seance_{date_exam.strftime('%Y%m%d')}_{seance.upper()}_{datetime.now().strftime('%H%M%S')}.docx"
        filepath = os.path.join(self.export_dir, filename)
        
        # Créer un document
        doc = Document()
        
        # Ajouter l'en-tête
        self._ajouter_entete(doc)
        
        # Ajouter le pied de page
        self._ajouter_pied_de_page(doc)
        
        # Ajouter le contenu de la séance (utiliser les strings pour l'affichage)
        self._ajouter_seance_au_document(doc, date_exam, h_debut_str, h_fin_str, examens)
        
        # Sauvegarder le document
        doc.save(filepath)
        
        logger.info(f"✅ Liste générée pour la séance {seance.upper()} du {date_exam}: {filepath}")
        return filepath

    # ===== NOUVELLES MÉTHODES POUR EXPORT PDF =====
    
    def _convertir_docx_vers_pdf(self, docx_path: str) -> str:
        """Convertit un fichier DOCX en PDF
        
        Args:
            docx_path: Chemin du fichier DOCX
            
        Returns:
            Chemin du fichier PDF généré
        """
        pdf_path = docx_path.replace('.docx', '.pdf')
        try:
            convert(docx_path, pdf_path)
            # Supprimer le fichier DOCX temporaire
            if os.path.exists(docx_path):
                os.remove(docx_path)
            return pdf_path
        except Exception as e:
            logger.error(f"Erreur lors de la conversion DOCX -> PDF: {str(e)}")
            raise
    
    def generer_convocations_individuelles_pdf(self) -> List[str]:
        """Génère les convocations individuelles en PDF pour chaque enseignant"""
        filepaths = []
        
        enseignants = self.db.query(Enseignant).all()
        
        for enseignant in enseignants:
            # Récupérer les affectations
            affectations = self.db.query(Affectation).options(
                joinedload(Affectation.examen)
            ).filter(
                Affectation.enseignant_id == enseignant.id
            ).join(Examen).order_by(Examen.dateExam, Examen.h_debut).all()
            
            if not affectations:
                continue
            
            # Générer DOCX d'abord
            filename_docx = f"convocation_{enseignant.nom}_{enseignant.prenom}_{datetime.now().strftime('%Y%m%d')}.docx"
            filepath_docx = os.path.join(self.export_dir, filename_docx)
            self._generer_convocation_word(enseignant, affectations, filepath_docx)
            
            # Convertir en PDF
            filepath_pdf = self._convertir_docx_vers_pdf(filepath_docx)
            filepaths.append(filepath_pdf)
        
        logger.info(f"✅ {len(filepaths)} convocations PDF générées")
        return filepaths
    

    
    def generer_listes_par_creneau_pdf(self) -> List[str]:
        """Génère les listes de surveillants par créneau en PDF - Un fichier par jour"""
        filepaths_pdf = []
        
        # Générer d'abord les fichiers DOCX
        filepaths_docx = self.generer_listes_par_creneau()
        
        # Convertir chaque DOCX en PDF
        for docx_path in filepaths_docx:
            try:
                pdf_path = self._convertir_docx_vers_pdf(docx_path)
                filepaths_pdf.append(pdf_path)
            except Exception as e:
                logger.error(f"Erreur lors de la conversion de {docx_path}: {str(e)}")
        
        logger.info(f"✅ {len(filepaths_pdf)} fichiers PDF générés (un par jour)")
        return filepaths_pdf
    

    
    def generer_convocation_enseignant_pdf(self, enseignant_id: int) -> str:
        """Génère la convocation PDF pour un enseignant spécifique"""
        # Générer d'abord le DOCX
        filepath_docx = self.generer_convocation_enseignant(enseignant_id)
        
        # Convertir en PDF
        filepath_pdf = self._convertir_docx_vers_pdf(filepath_docx)
        
        logger.info(f"✅ Convocation PDF générée: {filepath_pdf}")
        return filepath_pdf
    
    def generer_liste_creneau_specifique_pdf(self, date_exam: date, seance: str) -> str:
        """Génère la liste PDF des surveillants pour un créneau spécifique"""
        # Générer d'abord le DOCX
        filepath_docx = self.generer_liste_creneau_specifique(date_exam, seance)
        
        # Convertir en PDF
        filepath_pdf = self._convertir_docx_vers_pdf(filepath_docx)
        
        logger.info(f"✅ Liste PDF générée pour la séance {seance.upper()} du {date_exam}: {filepath_pdf}")
        return filepath_pdf
